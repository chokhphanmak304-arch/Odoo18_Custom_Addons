from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_background(cell, fill_color):
    """ตั้งค่าสีพื้นหลังของเซลล์"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_colored_box(doc, text, color, width=None):
    """เพิ่ม box ที่มีสีพื้นหลัง"""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    if width:
        table.width = width
    
    cell = table.rows[0].cells[0]
    cell.text = text
    set_cell_background(cell, color)
    
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(255, 255, 255)
    
    return table

def add_section_header(doc, text, color='2E5090'):
    """เพิ่มหัวข้อ section"""
    heading = doc.add_heading(text, level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # เพิ่มเส้นใต้สีสัน
    for run in heading.runs:
        run.font.color.rgb = RGBColor(46, 80, 144)
        run.font.size = Pt(14)
        run.font.bold = True

def create_detailed_workflow():
    """สร้าง Workflow Document แบบละเอียด"""
    doc = Document()
    
    # ตั้งค่าขอบหน้า
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # ============= หน้าปกเอกสาร =============
    title = doc.add_heading('TRANSPORT BOOKING MODULE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 60, 120)
    
    subtitle = doc.add_heading('Workflow Process Documentation', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # ข้อมูลเอกสาร
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ('Module Name', 'Transport Booking'),
        ('Version', '1.0'),
        ('Created Date', datetime.now().strftime('%Y-%m-%d')),
        ('Documentation Type', 'Workflow Process')
    ]
    
    for idx, (label, value) in enumerate(info_data):
        info_table.rows[idx].cells[0].text = label
        info_table.rows[idx].cells[1].text = value
        info_table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_background(info_table.rows[idx].cells[0], 'D9E8F5')
    
    doc.add_page_break()
    
    # ============= Table of Contents =============
    toc = doc.add_heading('สารบัญ (Table of Contents)', level=1)
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        '1. Business Overview',
        '2. Workflow State Diagram',
        '3. Detailed State Descriptions',
        '4. State Transition Rules',
        '5. Field Changes by State',
        '6. User Permissions & Access Control',
        '7. Business Logic & Validation',
        '8. Error Handling',
        '9. Integration Points',
        '10. System Flow'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ============= Business Overview =============
    add_section_header(doc, '1. Business Overview')
    
    overview_text = """ระบบ Transport Booking เป็นโมดูลสำหรับการจัดการการจองการเดินทาง โดยมีขั้นตอนการจัดการจากการสร้างจอง ยืนยัน ชำระเงิน มอบหมายคนขับ ไปจนถึงการสิ้นสุดการเดินทาง"""
    doc.add_paragraph(overview_text)
    
    doc.add_heading('วัตถุประสงค์หลัก (Main Objectives):', level=3)
    objectives = [
        'จัดการข้อมูลการจองการเดินทาง',
        'ติดตามสถานะของการจองอย่างเรียลไทม์',
        'จัดการการชำระเงินและยืนยัน',
        'มอบหมายคนขับและยานพาหนะ',
        'บันทึกข้อมูลการเดินทาง',
        'สร้างรายงานและสถิติ'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_page_break()
    
    # ============= Workflow State Diagram =============
    add_section_header(doc, '2. Workflow State Diagram')
    
    doc.add_paragraph('แผนภาพสถานะของการไหลของการจองการเดินทาง:')
    doc.add_paragraph()
    
    # วาด flow diagram แบบข้อความ
    flow_diagram = """
    ┌──────────────────────────────────────────────────────────────────────────┐
    │                          TRANSPORT BOOKING WORKFLOW                      │
    └──────────────────────────────────────────────────────────────────────────┘
    
    ┌─────────┐
    │ DRAFT   │  ← ผู้ใช้สร้างจองใหม่
    └────┬────┘
         │ (User Confirm)
         ▼
    ┌──────────────┐
    │ CONFIRMED    │  ← ยืนยันรายละเอียด
    └────┬─────────┘
         │ (Process Payment)
         ▼
    ┌──────────────┐
    │ PAID         │  ← ชำระเงินสำเร็จ
    └────┬─────────┘
         │ (Assign Driver)
         ▼
    ┌──────────────┐
    │ ASSIGNED     │  ← มอบหมายคนขับ
    └────┬─────────┘
         │ (Start Trip)
         ▼
    ┌──────────────┐
    │ IN PROGRESS  │  ← กำลังเดินทาง
    └────┬─────────┘
         │ (Complete Trip)
         ▼
    ┌──────────────┐
    │ COMPLETED    │  ← เดินทางสำเร็จ
    └────┬─────────┘
         │ (Close Booking)
         ▼
    ┌──────────────┐
    │ CLOSED       │  ← ปิดจองเรียบร้อย
    └──────────────┘
    
    
    ┌──────────────────────────────────────────────────────────────┐
    │ Exception Path:  Any State ──→ [CANCEL] ──→ CANCELLED       │
    │                  (ยกเลิกการจอง)                              │
    └──────────────────────────────────────────────────────────────┘
    """
    
    flow_para = doc.add_paragraph(flow_diagram)
    flow_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in flow_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ============= Detailed State Descriptions =============
    add_section_header(doc, '3. Detailed State Descriptions')
    
    states_detail = [
        {
            'name': 'DRAFT',
            'color': 'E8E8E8',
            'icon': '📋',
            'description': 'สถานะเริ่มต้น - จองใหม่ที่ยังไม่ยืนยัน',
            'details': [
                'ผู้ใช้สร้างจองใหม่',
                'สามารถแก้ไขข้อมูลได้ทั้งหมด',
                'ยังไม่มีการใช้ทรัพยากร',
                'สามารถลบได้'
            ],
            'fields_editable': ['Passenger Name', 'Pickup Location', 'Dropoff Location', 'Trip Date', 'Notes']
        },
        {
            'name': 'CONFIRMED',
            'color': 'B4C7E7',
            'icon': '✓',
            'description': 'จองได้รับการยืนยัน - รายละเอียดเสร็จสมบูรณ์',
            'details': [
                'ผู้ใช้ยืนยันรายละเอียด',
                'ทำการ validate ข้อมูล',
                'สร้างอ้างอิงจอง',
                'เตรียมบิลค่าบริการ',
                'ไม่สามารถแก้ไข passenger info ได้'
            ],
            'fields_editable': ['Notes', 'Special Requests']
        },
        {
            'name': 'PAID',
            'color': 'C6E0B4',
            'icon': '💳',
            'description': 'ชำระเงินสำเร็จ - พร้อมมอบหมายคนขับ',
            'details': [
                'ระบบประมวลผลการชำระเงิน',
                'บันทึก payment record',
                'ส่งหลักฐาน payment',
                'พร้อมสำหรับการมอบหมาย',
                'เปิดให้มอบหมายคนขับ'
            ],
            'fields_editable': ['Driver Notes']
        },
        {
            'name': 'ASSIGNED',
            'color': 'F4B084',
            'icon': '👤',
            'description': 'มอบหมายคนขับแล้ว - รอให้คนขับเริ่มเดินทาง',
            'details': [
                'ระบบมอบหมายคนขับและยานพาหนะ',
                'ส่ง notification ให้คนขับ',
                'บันทึกข้อมูลคนขับและยาน',
                'รอการเริ่มต้นการเดินทาง',
                'ผู้ใช้ได้รับ driver info'
            ],
            'fields_editable': ['Special Instructions']
        },
        {
            'name': 'IN PROGRESS',
            'color': 'FFE699',
            'icon': '🚗',
            'description': 'อยู่ระหว่างการเดินทาง - GPS กำลังติดตาม',
            'details': [
                'คนขับเริ่มเดินทาง',
                'บันทึก start time และ location',
                'เปิดใช้ GPS tracking',
                'ติดตามทำนายเวลาถึง',
                'ส่ง notification อัปเดตให้ผู้ใช้',
                'ผู้ใช้สามารถดูตำแหน่งคนขับ'
            ],
            'fields_editable': []
        },
        {
            'name': 'COMPLETED',
            'color': 'C6E0B4',
            'icon': '✅',
            'description': 'เดินทางสำเร็จ - รอการปิดจอง',
            'details': [
                'คนขับอัพเดท end time',
                'บันทึก end location',
                'หยุด GPS tracking',
                'คำนวณค่าใช้จ่ายสุดท้าย',
                'สร้างใบเสร็จ',
                'รอ confirmation ปิดจอง'
            ],
            'fields_editable': ['Completion Notes', 'Rating']
        },
        {
            'name': 'CLOSED',
            'color': '92D050',
            'icon': '🏁',
            'description': 'ปิดจองเรียบร้อย - เสร็จสิ้น',
            'details': [
                'จองปิดสิ้นสุด',
                'บันทึก final report',
                'สร้างสถิติ',
                'ไม่สามารถแก้ไขได้',
                'เก็บเอกสารประวัติ'
            ],
            'fields_editable': []
        },
        {
            'name': 'CANCELLED',
            'color': 'F8CBAD',
            'icon': '❌',
            'description': 'ยกเลิกการจอง - ยกเลิกจากการจอง',
            'details': [
                'ผู้ใช้หรือแอดมิน ยกเลิก',
                'บันทึกเหตุผลการยกเลิก',
                'คืน payment (ถ้ามี)',
                'แจ้งคนขับ (ถ้า assigned)',
                'ปิดจอง'
            ],
            'fields_editable': ['Cancellation Reason']
        }
    ]
    
    for state in states_detail:
        # Heading
        state_heading = doc.add_heading(f"{state['icon']} {state['name']}", level=3)
        for run in state_heading.runs:
            run.font.color.rgb = RGBColor(30, 60, 120)
        
        # Description
        desc_para = doc.add_paragraph(state['description'])
        desc_para.runs[0].font.italic = True
        
        # Details table
        table = doc.add_table(rows=len(state['details']) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header
        table.rows[0].cells[0].text = 'รายละเอียด'
        table.rows[0].cells[1].text = 'คำอธิบาย'
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(cell, state['color'])
        
        # Details
        for idx, detail in enumerate(state['details'], 1):
            table.rows[idx].cells[0].text = f"{idx}."
            table.rows[idx].cells[1].text = detail
        
        # Editable Fields
        if state['fields_editable']:
            doc.add_paragraph('สามารถแก้ไขได้: ' + ', '.join(state['fields_editable']), style='List Bullet')
        else:
            doc.add_paragraph('ไม่สามารถแก้ไขข้อมูลได้', style='List Bullet')
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= State Transition Rules =============
    add_section_header(doc, '4. State Transition Rules')
    
    transitions = [
        {
            'from': 'DRAFT',
            'to': 'CONFIRMED',
            'action': 'action_confirm()',
            'permission': 'User',
            'condition': [
                'ข้อมูลผู้โดยสารครบถ้วน',
                'เลือกสถานที่จอย-ลงแล้ว',
                'เลือกวันเดินทาง'
            ],
            'required_fields': ['passenger_name', 'phone', 'pickup_location', 'dropoff_location', 'trip_date']
        },
        {
            'from': 'CONFIRMED',
            'to': 'PAID',
            'action': 'action_pay()',
            'permission': 'User/Payment System',
            'condition': [
                'ชำระเงินผ่านระบบ',
                'สถานะ payment = Success',
                'ยืนยัน transaction'
            ],
            'required_fields': ['amount', 'payment_method']
        },
        {
            'from': 'PAID',
            'to': 'ASSIGNED',
            'action': 'action_assign_driver()',
            'permission': 'Admin/System',
            'condition': [
                'มีคนขับพร้อม',
                'มียานพาหนะพร้อม',
                'ตรวจสอบระยะทาง'
            ],
            'required_fields': ['driver_id', 'vehicle_id']
        },
        {
            'from': 'ASSIGNED',
            'to': 'IN PROGRESS',
            'action': 'action_start_trip()',
            'permission': 'Driver',
            'condition': [
                'คนขับยืนยันเริ่มเดินทาง',
                'GPS พร้อม',
                'เปิด tracking'
            ],
            'required_fields': ['start_time', 'start_location']
        },
        {
            'from': 'IN PROGRESS',
            'to': 'COMPLETED',
            'action': 'action_complete()',
            'permission': 'Driver',
            'condition': [
                'ถึงปลายทาง',
                'คนขับยืนยันเสร็จสิ้น',
                'บันทึก location'
            ],
            'required_fields': ['end_time', 'end_location']
        },
        {
            'from': 'COMPLETED',
            'to': 'CLOSED',
            'action': 'action_close()',
            'permission': 'User/Admin',
            'condition': [
                'ผู้โดยสารยืนยัน',
                'ไม่มี dispute',
                'คำนวณค่าใช้จ่ายเสร็จ'
            ],
            'required_fields': []
        },
        {
            'from': 'ANY',
            'to': 'CANCELLED',
            'action': 'action_cancel()',
            'permission': 'User/Admin',
            'condition': [
                'ผู้ใช้ร้องขอยกเลิก',
                'อยู่ในเวลาอนุญาต',
                'บันทึกเหตุผล'
            ],
            'required_fields': ['cancellation_reason']
        }
    ]
    
    for trans in transitions:
        trans_table = doc.add_table(rows=7, cols=2)
        trans_table.style = 'Light Grid'
        
        trans_data = [
            ('From State', trans['from']),
            ('To State', trans['to']),
            ('Action Method', trans['action']),
            ('Required Permission', trans['permission']),
            ('Conditions', '\n'.join(trans['condition'])),
            ('Required Fields', ', '.join(trans['required_fields']) if trans['required_fields'] else 'None'),
        ]
        
        for idx, (label, value) in enumerate(trans_data):
            trans_table.rows[idx].cells[0].text = label
            trans_table.rows[idx].cells[1].text = value
            trans_table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
            set_cell_background(trans_table.rows[idx].cells[0], 'D9E8F5')
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= Field Changes by State =============
    add_section_header(doc, '5. Field Changes by State')
    
    field_changes = {
        'DRAFT': {'editable': 'All', 'visible': 'All', 'locked': 'None'},
        'CONFIRMED': {'editable': 'Notes, Special Requests', 'visible': 'All', 'locked': 'Passenger Info'},
        'PAID': {'editable': 'Driver Notes', 'visible': 'All', 'locked': 'All Payment Fields'},
        'ASSIGNED': {'editable': 'Special Instructions', 'visible': 'All', 'locked': 'Driver & Vehicle Info'},
        'IN PROGRESS': {'editable': 'None', 'visible': 'All', 'locked': 'All (except logs)'},
        'COMPLETED': {'editable': 'Completion Notes, Rating', 'visible': 'All', 'locked': 'Trip Data'},
        'CLOSED': {'editable': 'None', 'visible': 'All + Reports', 'locked': 'All'},
        'CANCELLED': {'editable': 'Cancellation Reason', 'visible': 'Limited', 'locked': 'Most Fields'}
    }
    
    field_table = doc.add_table(rows=len(field_changes) + 1, cols=4)
    field_table.style = 'Light Grid Accent 1'
    
    headers = field_table.rows[0].cells
    headers[0].text = 'State'
    headers[1].text = 'Editable Fields'
    headers[2].text = 'Visible Fields'
    headers[3].text = 'Locked Fields'
    
    for cell in headers:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '2E5090')
    
    for idx, (state, changes) in enumerate(field_changes.items(), 1):
        field_table.rows[idx].cells[0].text = state
        field_table.rows[idx].cells[1].text = changes['editable']
        field_table.rows[idx].cells[2].text = changes['visible']
        field_table.rows[idx].cells[3].text = changes['locked']
    
    doc.add_page_break()
    
    # ============= User Permissions =============
    add_section_header(doc, '6. User Permissions & Access Control')
    
    doc.add_heading('Permission Matrix:', level=3)
    
    permissions = {
        'User': ['View Own Bookings', 'Create Booking', 'Confirm Booking', 'Pay', 'View Tracking', 'Cancel (before assigned)', 'Rate & Review'],
        'Driver': ['View Assigned Bookings', 'Start Trip', 'Complete Trip', 'View Route'],
        'Admin': ['View All Bookings', 'Assign Driver', 'Override Status', 'Process Payments', 'Manage Disputes', 'Generate Reports']
    }
    
    for role, perms in permissions.items():
        doc.add_heading(f'📍 {role}:', level=4)
        for perm in perms:
            doc.add_paragraph(perm, style='List Bullet')
    
    doc.add_page_break()
    
    # ============= Business Logic & Validation =============
    add_section_header(doc, '7. Business Logic & Validation')
    
    validations = [
        ('Email Validation', 'Email format ต้อง valid'),
        ('Phone Number', 'เบอร์โทรต้องเป็นหมายเลข valid'),
        ('Location Validation', 'ต้องเลือก pickup และ dropoff ต่างกัน'),
        ('Trip Date', 'วันจองต้องไม่น้อยกว่าวันปัจจุบัน 1 ชั่วโมง'),
        ('Payment Amount', 'ยอดเงินต้องมากกว่า 0'),
        ('Distance Calculation', 'คำนวณระยะทางจาก location 2 จุด'),
        ('Price Calculation', 'Base Price + (Distance × Rate per km)'),
        ('Driver Availability', 'ตรวจสอบคนขับว่าไม่มี conflict')
    ]
    
    val_table = doc.add_table(rows=len(validations) + 1, cols=2)
    val_table.style = 'Light Grid Accent 1'
    
    val_table.rows[0].cells[0].text = 'Validation Type'
    val_table.rows[0].cells[1].text = 'Rule'
    for cell in val_table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '2E5090')
    
    for idx, (val_type, rule) in enumerate(validations, 1):
        val_table.rows[idx].cells[0].text = val_type
        val_table.rows[idx].cells[1].text = rule
    
    doc.add_page_break()
    
    # ============= Error Handling =============
    add_section_header(doc, '8. Error Handling')
    
    errors = [
        {
            'error': 'Insufficient Payment',
            'scenario': 'ชำระเงินน้อยกว่า Total Amount',
            'handling': 'ปฏิเสธ transition ไป PAID'
        },
        {
            'error': 'No Driver Available',
            'scenario': 'ไม่มีคนขับพร้อมในเวลาที่กำหนด',
            'handling': 'แจ้งเตือนผู้ใช้ ขอเปลี่ยนเวลา'
        },
        {
            'error': 'Expired Booking',
            'scenario': 'จองเกินเวลา 30 นาทีของการเริ่มต้น',
            'handling': 'Auto-cancel และ refund'
        },
        {
            'error': 'Invalid Location',
            'scenario': 'Location ไม่พบในระบบ',
            'handling': 'ขอให้ผู้ใช้ระบุ coordinates'
        },
        {
            'error': 'GPS Timeout',
            'scenario': 'GPS ไม่ response เกิน 5 นาที',
            'handling': 'บันทึก log, ส่ง alert'
        }
    ]
    
    for error in errors:
        err_table = doc.add_table(rows=4, cols=2)
        err_table.style = 'Light Grid'
        
        err_data = [
            ('Error Type', error['error']),
            ('Scenario', error['scenario']),
            ('Handling', error['handling'])
        ]
        
        for idx, (label, value) in enumerate(err_data):
            err_table.rows[idx].cells[0].text = label
            err_table.rows[idx].cells[1].text = value
            err_table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
            set_cell_background(err_table.rows[idx].cells[0], 'F4B084')
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= Integration Points =============
    add_section_header(doc, '9. Integration Points')
    
    integrations = [
        {
            'name': 'Payment Gateway',
            'models': 'account.payment, account.invoice',
            'action': 'Process payment at CONFIRMED → PAID',
            'api': 'Payment API Integration'
        },
        {
            'name': 'GPS Tracking',
            'models': 'tracking.location',
            'action': 'Track location in IN PROGRESS state',
            'api': 'Google Maps API or similar'
        },
        {
            'name': 'Notifications',
            'models': 'mail.message, ir.actions.act_window',
            'action': 'Send email/SMS at state changes',
            'api': 'SMS Gateway, Email Server'
        },
        {
            'name': 'Reporting',
            'models': 'ir.actions.report',
            'action': 'Generate reports at COMPLETED/CLOSED',
            'api': 'Report Engine'
        }
    ]
    
    int_table = doc.add_table(rows=len(integrations) + 1, cols=4)
    int_table.style = 'Light Grid Accent 1'
    
    int_headers = int_table.rows[0].cells
    int_headers[0].text = 'Integration'
    int_headers[1].text = 'Related Models'
    int_headers[2].text = 'Trigger Action'
    int_headers[3].text = 'External API'
    
    for cell in int_headers:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '2E5090')
    
    for idx, integ in enumerate(integrations, 1):
        int_table.rows[idx].cells[0].text = integ['name']
        int_table.rows[idx].cells[1].text = integ['models']
        int_table.rows[idx].cells[2].text = integ['action']
        int_table.rows[idx].cells[3].text = integ['api']
    
    doc.add_page_break()
    
    # ============= System Flow =============
    add_section_header(doc, '10. System Flow - Complete Journey')
    
    flow_steps = [
        '1. Passenger creates booking (DRAFT)',
        '2. System validates passenger data',
        '3. Passenger confirms booking (CONFIRMED)',
        '4. System calculates fare and creates invoice',
        '5. Passenger makes payment (PAID)',
        '6. System processes payment via Payment Gateway',
        '7. Admin/System assigns driver (ASSIGNED)',
        '8. System sends notification to driver',
        '9. Driver accepts and starts trip (IN PROGRESS)',
        '10. System activates GPS tracking',
        '11. Real-time location updates sent to passenger',
        '12. Driver reaches destination (COMPLETED)',
        '13. Passenger confirms arrival and rates trip',
        '14. System generates receipt and booking closed (CLOSED)',
        '15. System records completion in analytics'
    ]
    
    for step in flow_steps:
        para = doc.add_paragraph(step)
        if '→' in step or 'System' in step:
            para.paragraph_format.left_indent = Inches(0.25)
        else:
            para.paragraph_format.left_indent = Inches(0)
    
    # Footer
    doc.add_page_break()
    footer_section = doc.add_heading('Documentation End', level=2)
    footer_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # บันทึกไฟล์
    output_path = r'C:\Program Files\Odoo 18.0.20251009\server\custom-addons\transport_booking\Workflow_Diagram_Detailed.docx'
    doc.save(output_path)
    print(f"✓ Detailed Workflow Diagram created: {output_path}")
    return output_path

if __name__ == '__main__':
    create_detailed_workflow()
