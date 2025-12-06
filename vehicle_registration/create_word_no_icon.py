# -*- coding: utf-8 -*-
"""
สร้างเอกสาร Word - กระบวนการทำงาน Vehicle Registration Module (ไม่มี Icon)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading_with_color(doc, text, level, color_rgb):
    """เพิ่มหัวเรื่องพร้อมสี"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

# สร้างเอกสาร
doc = Document()

# ตั้งชื่อเอกสาร
title = doc.add_heading('ระบบจัดการรถและการซ่อมบำรุง', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Odoo 18.0 - กระบวนการทำงาน (Workflow)', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ข้อมูลโมดูล
info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Light Grid Accent 1'

info_data = [
    ('Module:', 'Vehicle Registration & Maintenance'),
    ('Version:', '18.0.1.1.0'),
    ('ประเภท:', 'Operations/Inventory'),
    ('วันที่สร้าง:', '27 ตุลาคม 2568'),
    ('ที่อยู่โฟลเดอร์:', 'custom-addons/vehicle_registration/'),
]

for i, (key, value) in enumerate(info_data):
    info_table.rows[i].cells[0].text = key
    info_table.rows[i].cells[1].text = value
    info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

# ===== 1. บทนำ =====
add_heading_with_color(doc, '1. บทนำ', 1, (46, 125, 50))

doc.add_paragraph('ระบบจัดการรถและการซ่อมบำรุง เป็นโมดูล Odoo ที่ออกแบบมาเพื่อจัดการ:')

intro_points = [
    'รถบรรทุก (ทั้งรถของบริษัท รถร่วม และรถเช่า)',
    'ผู้ขับขี่และใบขับขี่',
    'เอกสารรถ (ภาษี พ.ร.บ. ประกันภัย)',
    'การแจ้งซ่อมและประวัติการซ่อมบำรุง',
    'แจ้งเตือนอัตโนมัติภายในระบบ Odoo',
]

for point in intro_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()

# ===== 2. กระบวนการแจ้งและซ่อมรถ =====
add_heading_with_color(doc, '2. กระบวนการแจ้งและซ่อมรถ', 1, (25, 118, 210))

steps_maintenance = [
    ('ขั้นตอนที่ 1: ผู้ใช้แจ้งปัญหา', 'ร่างแบบ (Draft)', 
     ['เลือกรถ (ทะเบียนรถ)', 'เลือกหมวดหมู่ปัญหา (เครื่องยนต์/ระบบเกียร์/ระบบเบรก/ระบบไฟฟ้า/ยาง/ตัวถัง/อื่นๆ)', 'เขียนรายละเอียดปัญหา', 'ระบุความเร่งด่วน (ปกติ/ด่วน/ด่วนมาก)'],
     'ระบบสร้างรหัสอัตโนมัติ (MR-001, MR-002...)'),
    
    ('ขั้นตอนที่ 2: ผู้บริหารยืนยันการแจ้ง', 'รอการซ่อม (Pending)',
     ['คลิก "ยืนยัน"', 'มอบหมายให้ช่าง'],
     'สถานะบันทึก: draft → pending | สถานะรถ: available → maintenance | ส่งแจ้งเตือน'),
    
    ('ขั้นตอนที่ 3: ช่างเริ่มต้นซ่อม', 'กำลังซ่อม (In Progress)',
     ['คลิก "เริ่มซ่อม"', 'บันทึกเวลาเริ่ม'],
     'สถานะบันทึก: pending → in_progress | เวลาเริ่มซ่อม: บันทึก | ช่างสามารถป้อนค่าใช้จ่าย'),
    
    ('ขั้นตอนที่ 4: ช่างแจ้งว่าซ่อมเสร็จ', 'เสร็จสิ้น (Done)',
     ['ป้อนค่าใช้จ่ายจริง (ข้อบังคับ)', 'คลิก "ซ่อมเสร็จ"'],
     'สร้างบันทึกประวัติการซ่อม | สถานะ: done | สถานะรถ: maintenance → available | ส่งแจ้งเตือน'),
    
    ('ขั้นตอนที่ 5: บันทึกประวัติการซ่อม', 'บันทึกอัตโนมัติ',
     ['รหัส: MH-001, MH-002...', 'ข้อมูลทั้งหมดจากการแจ้งซ่อม', 'สามารถสร้างแจ้งเตือนซ้ำได้'],
     'เก็บไว้สำหรับอ้างอิงในอนาคต'),
]

for title, status, actions, system in steps_maintenance:
    step_heading = doc.add_paragraph()
    step_heading.paragraph_format.left_indent = Inches(0.25)
    run = step_heading.add_run(title)
    run.font.bold = True
    run.font.size = Pt(12)
    
    status_p = doc.add_paragraph(f"สถานะ: {status}", style='List Bullet')
    status_p.paragraph_format.left_indent = Inches(0.5)
    
    for action in actions:
        action_p = doc.add_paragraph(action, style='List Bullet 2')
        action_p.paragraph_format.left_indent = Inches(0.75)
    
    system_p = doc.add_paragraph(f"ระบบ: {system}")
    system_p.paragraph_format.left_indent = Inches(0.5)
    system_p.runs[0].font.italic = True
    
    doc.add_paragraph()

# ===== 3. กระบวนการแจ้งเตือนแบบอัตโนมัติ =====
add_heading_with_color(doc, '3. กระบวนการแจ้งเตือนแบบอัตโนมัติ', 1, (123, 31, 162))

notification_steps = [
    ('ขั้นตอนที่ 1: สร้างแจ้งเตือน', 'ผู้จัดการเลือกบันทึกการซ่อม เปิดการแจ้งเตือน ตั้งช่วงวัน เลือกผู้รับแจ้งเตือน ตั้งวันเริ่ม-สิ้นสุด'),
    ('ขั้นตอนที่ 2: ระบบเช็คเงื่อนไข', 'CRON Job ทำงานทุกวันเวลา 12:00 น. ตรวจสอบ: (1) การแจ้งเตือนเปิดใช้ (2) อยู่ในช่วงวันที่ (3) ผ่านไป N วัน (4) มีผู้รับแจ้ง'),
    ('ขั้นตอนที่ 3: ส่งแจ้งเตือน', 'หากเงื่อนไขครบ: สร้างบันทึกแจ้งเตือน สร้าง Activity ในหน้าแรก แจ้งเตือนใน Inbox บันทึกเวลาส่ง'),
    ('ขั้นตอนที่ 4: ผู้รับดูแจ้งเตือน', 'ผู้รับเข้ามา Odoo ดู Inbox คลิกดูรายละเอียด อ่านข้อมูล เพิ่มความเห็นได้'),
    ('ขั้นตอนที่ 5: ทำเครื่องหมายว่าอ่านแล้ว', 'ผู้รับคลิก "อ่านแล้ว" → ระบบอัปเดต: status sent→read, ลบ Activity, หลุดจาก Inbox'),
    ('ขั้นตอนที่ 6: รอการแจ้งเตือนครั้งถัดไป', 'รอ N วัน (ตามช่วงที่ตั้ง) จากนั้น CRON ทำงานอีกครั้ง วนซ้ำไปเรื่อยๆ จนสิ้นสุด'),
]

for title, details in notification_steps:
    step_heading = doc.add_paragraph()
    step_heading.paragraph_format.left_indent = Inches(0.25)
    run = step_heading.add_run(title)
    run.font.bold = True
    run.font.size = Pt(12)
    
    details_p = doc.add_paragraph(details)
    details_p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph()

# ===== 4. กระบวนการจัดเก็บเอกสาร =====
add_heading_with_color(doc, '4. กระบวนการจัดเก็บเอกสาร', 1, (230, 81, 0))

document_steps = [
    ('ขั้นตอนที่ 1: เพิ่มเอกสารของรถ', 'ผู้จัดการกรอก: ชื่อเอกสาร ประเภท (ภาษี/พ.ร.บ./ประกันภัย/ทะเบียน/ใบอนุญาต/อื่นๆ) เลขที่เอกสาร วันออก วันหมดอายุ'),
    ('ขั้นตอนที่ 2: ระบบคำนวณสถานะ', 'ตรวจสอบทุกวัน: ยังใช้ได้ (หมดอายุ > today+30) | ใกล้หมดอายุ (today ≤ หมดอายุ ≤ today+30) | หมดอายุแล้ว (หมดอายุ < today)'),
    ('ขั้นตอนที่ 3: ส่งแจ้งเตือน', 'เมื่ออยู่ในสถานะ "ใกล้หมดอายุ" เลือกผู้รับแจ้งเตือน ส่งอีเมล ระบุ: ชื่อเอกสาร ทะเบียนรถ วันหมดอายุ เหลือกี่วัน'),
    ('ขั้นตอนที่ 4: ผู้จัดการต่ออายุ', 'ติดต่อหน่วยงาน ต่ออายุเอกสาร รับเอกสารใหม่ บันทึกในระบบ: เพิ่มบันทึกใหม่ ใส่วันหมดอายุใหม่ เอกสารเก่าเป็นประวัติ'),
    ('ขั้นตอนที่ 5: สถานะกลับเป็นใช้ได้', 'สถานะ: ใช้ได้ | รถสามารถใช้งานได้ | ไม่มีการแจ้งเตือน | พร้อมสำหรับการขนส่ง'),
]

for title, details in document_steps:
    step_heading = doc.add_paragraph()
    step_heading.paragraph_format.left_indent = Inches(0.25)
    run = step_heading.add_run(title)
    run.font.bold = True
    run.font.size = Pt(12)
    
    details_p = doc.add_paragraph(details)
    details_p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph()

# ===== 5. กระบวนการจัดการผู้ขับขี่ =====
add_heading_with_color(doc, '5. กระบวนการจัดการผู้ขับขี่', 1, (104, 159, 56))

driver_steps = [
    ('ขั้นตอนที่ 1: สมัครผู้ขับขี่', 'ผู้จัดการกรอก: ชื่อ-นามสกุล PIN 6 หลัก (ห้ามซ้ำ) เบอร์โทรศัพท์ เลขบัตรประชาชน เลขใบขับขี่ ประเภทใบขับขี่ วันหมดอายุ'),
    ('ขั้นตอนที่ 2: ระบบตรวจสอบ PIN', 'Validation: (1) ต้องเป็นตัวเลขเท่านั้น (2) ต้องเป็น 6 หลักพอดี (3) ห้ามเหมือนผู้ขับคนอื่น (UNIQUE) | ถ้าผิด: PIN นี้ถูกใช้งาน หรือ PIN ต้องเป็น 6 หลัก'),
    ('ขั้นตอนที่ 3: ระบบคำนวณสถานะใบขับขี่', 'ตรวจสอบทุกวัน: ใช้ได้ (หมดอายุ > today+30) | ใกล้หมดอายุ (today ≤ หมดอายุ ≤ today+30) | หมดอายุแล้ว (หมดอายุ < today)'),
    ('ขั้นตอนที่ 4: ส่งแจ้งเตือน', 'เมื่อสถานะ "ใกล้หมดอายุ" เลือกผู้รับแจ้งเตือน ส่งอีเมล ระบุ: ชื่อผู้ขับขี่ เลขใบขับขี่ วันหมดอายุ เหลือกี่วัน'),
    ('ขั้นตอนที่ 5: ผู้ขับขี่ต่อใบขับขี่', 'ไปต่ออายุที่จังหวัร รับใบขับขี่ใหม่ บอกให้ผู้จัดการทราบ → ระบบอัปเดตวันหมดอายุใหม่ → สถานะกลับ ใช้ได้'),
    ('ขั้นตอนที่ 6: ประวัติอุบัติเหตุ', 'บันทึกหากมีเหตุการณ์: วันเวลา สถานที่ รายละเอียด ความรุนแรง (เล็ก/ปาน/ร้ายแรง) ค่าเสียหาย → ติดตามประวัติผู้ขับ ประเมินความปลอดภัย'),
]

for title, details in driver_steps:
    step_heading = doc.add_paragraph()
    step_heading.paragraph_format.left_indent = Inches(0.25)
    run = step_heading.add_run(title)
    run.font.bold = True
    run.font.size = Pt(12)
    
    details_p = doc.add_paragraph(details)
    details_p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph()

# ===== สรุป =====
doc.add_page_break()
add_heading_with_color(doc, '6. สรุปและข้อมูลติดต่อ', 1, (25, 118, 210))

summary_table = doc.add_table(rows=5, cols=2)
summary_table.style = 'Light Grid Accent 1'

summary_data = [
    ('กระบวนการทั้งหมด', '4 กระบวนการ'),
    ('ขั้นตอนรวม', '22 ขั้นตอน'),
    ('สถานะที่ติดตาม', '7 สถานะ (Draft/Pending/In Progress/Done/Cancelled/Valid/Expired)'),
    ('ระบบอัตโนมัติ', 'CRON Job ทุกวัน / Status Computation / Notification'),
    ('Odoo Version', '18.0.1.1.0'),
]

for i, (key, value) in enumerate(summary_data):
    summary_table.rows[i].cells[0].text = key
    summary_table.rows[i].cells[1].text = value
    summary_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()

contact_heading = doc.add_heading('ข้อมูลติดต่อ', level=2)
contact_heading.runs[0].font.color.rgb = RGBColor(25, 118, 210)

contact_info = [
    'Module: Vehicle Registration & Maintenance',
    'รหัส: vehicle_registration',
    'เวอร์ชัน: 18.0.1.1.0',
    'ผู้พัฒนา: Your Company',
    'ลิขสิทธิ์: LGPL-3',
    'โฟลเดอร์: custom-addons/vehicle_registration/',
]

for info in contact_info:
    doc.add_paragraph(info, style='List Bullet')

doc.add_paragraph()
footer = doc.add_paragraph('เอกสารฉบับนี้สร้างขึ้น: 27 ตุลาคม 2568')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(10)
footer.runs[0].font.italic = True
footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

# บันทึกไฟล์
output_path = r'C:\Program Files\Odoo 18.0.20251009\server\custom-addons\vehicle_registration\Workflow_Diagram_No_Icon.docx'
doc.save(output_path)

print('สร้างไฟล์ Word สำเร็จ!')
print('ไฟล์ได้แก้ไขให้ไม่มี Icon แล้ว')
print(f'ชื่อไฟล์: Workflow_Diagram_No_Icon.docx')
