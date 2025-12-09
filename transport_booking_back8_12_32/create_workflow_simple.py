from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill_color):
    """ตั้งค่าสีพื้นหลังของเซลล์"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_section_header(doc, text):
    """เพิ่มหัวข้อ section"""
    heading = doc.add_heading(text, level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(46, 80, 144)
        run.font.size = Pt(14)
        run.font.bold = True

def create_simple_workflow():
    """สร้าง Workflow Document แบบง่าย"""
    doc = Document()
    
    # ตั้งค่าขอบหน้า
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # ============= หัวเรื่อง =============
    title = doc.add_heading('TRANSPORT BOOKING', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 60, 120)
    
    subtitle = doc.add_heading('Workflow Process', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ============= Overview =============
    add_section_header(doc, '1. Overview')
    doc.add_paragraph('ระบบจัดการการจองการเดินทาง ที่ติดตามสถานะจากการสร้างจองจนถึงการสิ้นสุดการเดินทาง')
    doc.add_paragraph()
    
    # ============= Main Flow =============
    add_section_header(doc, '2. Main Workflow States')
    
    # สร้าง flow diagram แบบ ASCII art
    flow_diagram = """
┌─────────┐
│  DRAFT  │  ← ผู้ใช้สร้างจองใหม่
└────┬────┘
     │
     ▼
┌──────────────┐
│ CONFIRMED    │  ← ยืนยันรายละเอียด
└────┬─────────┘
     │
     ▼
┌──────────────┐
│ PAID         │  ← ชำระเงินเสร็จ
└────┬─────────┘
     │
     ▼
┌──────────────┐
│ ASSIGNED     │  ← มอบหมายคนขับ
└────┬─────────┘
     │
     ▼
┌──────────────┐
│ IN PROGRESS  │  ← กำลังเดินทาง
└────┬─────────┘
     │
     ▼
┌──────────────┐
│ COMPLETED    │  ← เดินทางสำเร็จ
└────┬─────────┘
     │
     ▼
┌──────────────┐
│ CLOSED       │  ← ปิดจองเรียบร้อย
└──────────────┘
    """
    
    flow_para = doc.add_paragraph(flow_diagram)
    flow_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in flow_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # ============= State Details =============
    add_section_header(doc, '3. State Details')
    
    states_data = [
        ('DRAFT', 'สถานะเริ่มต้น', ['สร้างจองใหม่', 'สามารถแก้ไขข้อมูลได้ทั้งหมด', 'ยังไม่ชำระเงิน']),
        ('CONFIRMED', 'ยืนยันแล้ว', ['ยืนยันรายละเอียด', 'เตรียมชำระเงิน', 'ไม่สามารถแก้ไข passenger info']),
        ('PAID', 'ชำระเงิน', ['ชำระเงินสำเร็จ', 'เตรียมมอบหมายคนขับ', 'บันทึก payment']),
        ('ASSIGNED', 'มอบหมาย', ['มอบหมายคนขับและยาน', 'ส่ง notification ให้คนขับ', 'รอการเริ่มเดินทาง']),
        ('IN PROGRESS', 'เดินทาง', ['คนขับเริ่มเดินทาง', 'เปิด GPS tracking', 'ส่ง update ให้ผู้ใช้']),
        ('COMPLETED', 'สำเร็จ', ['คนขับถึงปลายทาง', 'ผู้โดยสารยืนยัน', 'บันทึก end time']),
        ('CLOSED', 'ปิด', ['จองปิดเสร็จสิ้น', 'สร้างใบเสร็จ', 'เก็บประวัติ'])
    ]
    
    for state_name, state_desc, details in states_data:
        # State name and description
        para = doc.add_paragraph()
        run = para.add_run(state_name)
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(30, 60, 120)
        para.add_run(f' - {state_desc}')
        
        # Details
        for detail in details:
            doc.add_paragraph(detail, style='List Bullet')
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= Transition Rules =============
    add_section_header(doc, '4. Transition Rules')
    
    transitions_table = doc.add_table(rows=8, cols=4)
    transitions_table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = transitions_table.rows[0].cells
    headers = ['From', 'To', 'Condition', 'Who']
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        for run in header_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(header_cells[i], '2E5090')
    
    # Data
    transitions_data = [
        ('DRAFT', 'CONFIRMED', 'ข้อมูลครบถ้วน', 'User'),
        ('CONFIRMED', 'PAID', 'ชำระเงินสำเร็จ', 'System/User'),
        ('PAID', 'ASSIGNED', 'มีคนขับพร้อม', 'Admin/System'),
        ('ASSIGNED', 'IN PROGRESS', 'คนขับเริ่ม', 'Driver'),
        ('IN PROGRESS', 'COMPLETED', 'ถึงปลายทาง', 'Driver'),
        ('COMPLETED', 'CLOSED', 'ผู้ใช้ยืนยัน', 'User/Admin'),
        ('ANY STATE', 'CANCELLED', 'ผู้ใช้ยกเลิก', 'User/Admin')
    ]
    
    for idx, (from_state, to_state, condition, who) in enumerate(transitions_data, 1):
        transitions_table.rows[idx].cells[0].text = from_state
        transitions_table.rows[idx].cells[1].text = to_state
        transitions_table.rows[idx].cells[2].text = condition
        transitions_table.rows[idx].cells[3].text = who
    
    doc.add_page_break()
    
    # ============= Editable Fields =============
    add_section_header(doc, '5. Editable Fields by State')
    
    fields_data = [
        ('DRAFT', 'ทั้งหมด', 'Passenger name, Phone, Locations, Date, Time'),
        ('CONFIRMED', 'จำกัด', 'Notes, Special requests'),
        ('PAID', 'จำกัด', 'Driver notes'),
        ('ASSIGNED', 'จำกัด', 'Special instructions'),
        ('IN PROGRESS', 'ไม่มี', 'ไม่สามารถแก้ไข'),
        ('COMPLETED', 'จำกัด', 'Completion notes, Rating'),
        ('CLOSED', 'ไม่มี', 'ไม่สามารถแก้ไข'),
        ('CANCELLED', 'จำกัด', 'Cancellation reason')
    ]
    
    fields_table = doc.add_table(rows=len(fields_data) + 1, cols=3)
    fields_table.style = 'Light Grid Accent 1'
    
    fields_headers = fields_table.rows[0].cells
    fields_headers[0].text = 'State'
    fields_headers[1].text = 'Permission'
    fields_headers[2].text = 'Fields'
    
    for cell in fields_headers:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '2E5090')
    
    for idx, (state, permission, fields) in enumerate(fields_data, 1):
        fields_table.rows[idx].cells[0].text = state
        fields_table.rows[idx].cells[1].text = permission
        fields_table.rows[idx].cells[2].text = fields
    
    doc.add_page_break()
    
    # ============= User Roles =============
    add_section_header(doc, '6. User Roles & Permissions')
    
    roles_data = {
        'END USER (ผู้โดยสาร)': [
            'สร้างจองใหม่ (DRAFT)',
            'ยืนยันจอง (CONFIRMED)',
            'ชำระเงิน (PAID)',
            'ดูสถานะและตำแหน่ง (IN PROGRESS)',
            'ยืนยันสำเร็จ (COMPLETED)',
            'ให้คะแนน (CLOSED)',
            'ยกเลิก (ก่อน ASSIGNED)'
        ],
        'DRIVER (คนขับ)': [
            'ดูจองที่มอบหมาย (ASSIGNED)',
            'เริ่มเดินทาง (IN PROGRESS)',
            'อัพเดท GPS location',
            'สิ้นสุดการเดินทาง (COMPLETED)'
        ],
        'ADMIN': [
            'ดูจองทั้งหมด',
            'มอบหมายคนขับ (ASSIGNED)',
            'แก้ไขสถานะ',
            'จัดการการชำระเงิน',
            'ดูรายงาน'
        ]
    }
    
    for role, permissions in roles_data.items():
        para = doc.add_paragraph()
        run = para.add_run(role)
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(30, 60, 120)
        
        for perm in permissions:
            doc.add_paragraph(perm, style='List Bullet')
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= Key Actions =============
    add_section_header(doc, '7. Key Actions & Methods')
    
    actions_data = [
        ('action_confirm()', 'DRAFT → CONFIRMED', 'ผู้ใช้ยืนยันจอง'),
        ('action_pay()', 'CONFIRMED → PAID', 'ระบบประมวลผลการชำระเงิน'),
        ('action_assign_driver()', 'PAID → ASSIGNED', 'ระบบมอบหมายคนขับ'),
        ('action_start_trip()', 'ASSIGNED → IN PROGRESS', 'คนขับเริ่มเดินทาง'),
        ('action_complete()', 'IN PROGRESS → COMPLETED', 'คนขับสิ้นสุดเดินทาง'),
        ('action_close()', 'COMPLETED → CLOSED', 'ปิดจองเรียบร้อย'),
        ('action_cancel()', 'ANY → CANCELLED', 'ยกเลิกการจอง')
    ]
    
    actions_table = doc.add_table(rows=len(actions_data) + 1, cols=3)
    actions_table.style = 'Light Grid Accent 1'
    
    actions_headers = actions_table.rows[0].cells
    actions_headers[0].text = 'Method Name'
    actions_headers[1].text = 'Transition'
    actions_headers[2].text = 'Description'
    
    for cell in actions_headers:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '2E5090')
    
    for idx, (method, transition, description) in enumerate(actions_data, 1):
        actions_table.rows[idx].cells[0].text = method
        actions_table.rows[idx].cells[1].text = transition
        actions_table.rows[idx].cells[2].text = description
    
    doc.add_page_break()
    
    # ============= Database Fields =============
    add_section_header(doc, '8. Database Fields')
    
    fields_info = [
        ('name', 'Char', 'หมายเลขจอง'),
        ('state', 'Selection', 'สถานะปัจจุบัน'),
        ('passenger_name', 'Char', 'ชื่อผู้โดยสาร'),
        ('passenger_phone', 'Char', 'เบอร์โทรศัพท์'),
        ('passenger_email', 'Email', 'อีเมล'),
        ('pickup_location', 'Char', 'สถานที่รับ'),
        ('dropoff_location', 'Char', 'สถานที่ลง'),
        ('trip_date', 'Date', 'วันที่เดินทาง'),
        ('trip_time', 'Time', 'เวลาเดินทาง'),
        ('driver_id', 'Many2one', 'ชื่อคนขับ'),
        ('vehicle_id', 'Many2one', 'ยานพาหนะ'),
        ('amount', 'Float', 'ยอดเงิน'),
        ('payment_method', 'Selection', 'วิธีชำระเงิน'),
        ('payment_status', 'Selection', 'สถานะการชำระ'),
        ('start_time', 'Datetime', 'เวลาเริ่มเดินทาง'),
        ('start_location', 'Char', 'ตำแหน่งเริ่มต้น'),
        ('end_time', 'Datetime', 'เวลาสิ้นสุดเดินทาง'),
        ('end_location', 'Char', 'ตำแหน่งปลายทาง'),
        ('distance', 'Float', 'ระยะทาง (km)'),
        ('duration', 'Float', 'ระยะเวลา (นาที)'),
        ('rating', 'Integer', 'คะแนนการบริการ'),
        ('notes', 'Text', 'หมายเหตุ'),
        ('cancellation_reason', 'Text', 'เหตุผลการยกเลิก')
    ]
    
    db_table = doc.add_table(rows=len(fields_info) + 1, cols=3)
    db_table.style = 'Light Grid Accent 1'
    
    db_headers = db_table.rows[0].cells
    db_headers[0].text = 'Field Name'
    db_headers[1].text = 'Type'
    db_headers[2].text = 'Description'
    
    for cell in db_headers:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '2E5090')
    
    for idx, (field_name, field_type, description) in enumerate(fields_info, 1):
        db_table.rows[idx].cells[0].text = field_name
        db_table.rows[idx].cells[1].text = field_type
        db_table.rows[idx].cells[2].text = description
    
    doc.add_page_break()
    
    # ============= Validation Rules =============
    add_section_header(doc, '9. Validation Rules')
    
    validation_list = [
        'Passenger name ต้องไม่ว่าง',
        'Phone number ต้องเป็นหมายเลขที่ถูกต้อง',
        'Email ต้อง valid format',
        'Pickup location ≠ Dropoff location',
        'Trip date ต้องไม่น้อยกว่าวันปัจจุบัน 1 ชั่วโมง',
        'Amount ต้องมากกว่า 0',
        'Payment method ต้องเลือก',
        'Driver ต้องว่างในช่วงเวลาดังกล่าว',
        'Vehicle ต้องพร้อมใช้งาน'
    ]
    
    for rule in validation_list:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_page_break()
    
    # ============= Error Scenarios =============
    add_section_header(doc, '10. Error Scenarios')
    
    error_scenarios = [
        {
            'scenario': 'Insufficient Payment',
            'cause': 'ยอดเงินน้อยกว่า fare',
            'handling': 'ปฏิเสธการทำธุรกรรม'
        },
        {
            'scenario': 'No Driver Available',
            'cause': 'ไม่มีคนขับพร้อมในเวลาที่กำหนด',
            'handling': 'ส่ง notification ขอเปลี่ยนเวลา'
        },
        {
            'scenario': 'Trip Expired',
            'cause': 'จองเกินเวลา 30 นาทีของการเริ่มต้น',
            'handling': 'Auto-cancel และคืนเงิน'
        },
        {
            'scenario': 'GPS Tracking Failed',
            'cause': 'GPS ไม่ response',
            'handling': 'บันทึก log และส่ง alert'
        },
        {
            'scenario': 'Invalid Location',
            'cause': 'Location ไม่พบในระบบ',
            'handling': 'ขอให้ผู้ใช้ระบุพิกัด'
        }
    ]
    
    for error in error_scenarios:
        para = doc.add_paragraph()
        run = para.add_run(f"{error['scenario']}")
        run.font.bold = True
        para.add_run(f" - {error['cause']}")
        doc.add_paragraph(f"วิธีแก้ไข: {error['handling']}", style='List Bullet')
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ============= Integration Points =============
    add_section_header(doc, '11. Integration Points')
    
    integration_list = [
        ('Payment Gateway', 'ชำระเงินออนไลน์'),
        ('GPS System', 'ติดตามตำแหน่งคนขับ'),
        ('Email Service', 'ส่ง confirmation และ notification'),
        ('SMS Gateway', 'ส่ง SMS update'),
        ('Reporting System', 'สร้างรายงาน')
    ]
    
    int_table = doc.add_table(rows=len(integration_list) + 1, cols=2)
    int_table.style = 'Light Grid Accent 1'
    
    int_headers = int_table.rows[0].cells
    int_headers[0].text = 'Integration'
    int_headers[1].text = 'Purpose'
    
    for cell in int_headers:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '2E5090')
    
    for idx, (integration, purpose) in enumerate(integration_list, 1):
        int_table.rows[idx].cells[0].text = integration
        int_table.rows[idx].cells[1].text = purpose
    
    # Footer
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('End of Documentation')
    footer_run.font.italic = True
    footer_run.font.size = Pt(11)
    
    # บันทึกไฟล์
    output_path = r'C:\Program Files\Odoo 18.0.20251009\server\custom-addons\transport_booking\Workflow_Diagram.docx'
    doc.save(output_path)
    print(f"Success: {output_path}")
    return output_path

if __name__ == '__main__':
    create_simple_workflow()
