from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_for_workflow(doc, title, states):
    """สร้างตาราสำหรับ workflow"""
    # เพิ่มหัวข้อ
    heading = doc.add_heading(title, level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # เพิ่มตาราแสดง workflow states
    table = doc.add_table(rows=len(states) + 1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # หัวตาราง
    header_cells = table.rows[0].cells
    headers = ['ลำดับ', 'สถานะ (State)', 'คำอธิบาย', 'การกระทำถัดไป']
    
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        # จัดฟอร์แมตหัวตาราง
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # เพิ่มสีพื้นหลัง
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '366092')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # เพิ่มข้อมูลแต่ละ state
    for idx, state in enumerate(states, 1):
        row_cells = table.rows[idx].cells
        row_cells[0].text = str(idx)
        row_cells[1].text = state['state']
        row_cells[2].text = state['description']
        row_cells[3].text = state['next_action']
        
        # จัดตำแหน่งตัวอักษรกึ่งกลาง
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def create_transport_booking_workflow():
    """สร้าง Transport Booking Workflow Document"""
    doc = Document()
    
    # ตั้งค่าขอบหน้า
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # หัวเอกสาร
    title = doc.add_heading('Transport Booking - Workflow Diagram', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Odoo Custom Module')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(12)
    subtitle_format.font.italic = True
    
    doc.add_paragraph()  # เพิ่มว่าง
    
    # =================================================================
    # Booking Request Workflow
    # =================================================================
    booking_states = [
        {
            'state': 'Draft',
            'description': 'จองใหม่ยังไม่ยืนยัน',
            'next_action': 'ยืนยัน (Confirm) → Confirmed'
        },
        {
            'state': 'Confirmed',
            'description': 'จองได้รับการยืนยัน',
            'next_action': 'ชำระเงิน (Pay) → Paid'
        },
        {
            'state': 'Paid',
            'description': 'ชำระเงินสำเร็จ',
            'next_action': 'มอบหมายคนขับ (Assign) → Assigned'
        },
        {
            'state': 'Assigned',
            'description': 'มอบหมายคนขับแล้ว',
            'next_action': 'เริ่มการเดินทาง (Start) → In Progress'
        },
        {
            'state': 'In Progress',
            'description': 'อยู่ระหว่างการเดินทาง',
            'next_action': 'สิ้นสุดการเดินทาง (Complete) → Completed'
        },
        {
            'state': 'Completed',
            'description': 'เดินทางสำเร็จ',
            'next_action': 'ปิด (Close) → Closed'
        },
        {
            'state': 'Closed',
            'description': 'จองปิดสิ้นสุด',
            'next_action': 'เสร็จสิ้น'
        },
        {
            'state': 'Cancelled',
            'description': 'ยกเลิกการจอง',
            'next_action': 'เสร็จสิ้น'
        }
    ]
    
    add_table_for_workflow(doc, 'Booking Request States', booking_states)
    
    doc.add_page_break()
    
    # =================================================================
    # Workflow Flow Chart (ข้อความบรรยาย)
    # =================================================================
    flow_section = doc.add_heading('Workflow Flow', level=2)
    flow_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    flow_text = """
Draft (จองใหม่)
      ↓
Confirm (ยืนยัน)
      ↓
Confirmed (ยืนยันแล้ว)
      ↓
Pay (ชำระเงิน)
      ↓
Paid (ชำระสำเร็จ)
      ↓
Assign Driver (มอบหมายคนขับ)
      ↓
Assigned (มอบหมายแล้ว)
      ↓
Start Trip (เริ่มเดินทาง)
      ↓
In Progress (อยู่ระหว่างเดินทาง)
      ↓
Complete (สิ้นสุดเดินทาง)
      ↓
Completed (เดินทางสำเร็จ)
      ↓
Close (ปิด)
      ↓
Closed (ปิดสิ้นสุด)

[สามารถยกเลิก (Cancel) ได้ทุกขั้นตอน → Cancelled]
    """
    
    flow_para = doc.add_paragraph(flow_text)
    flow_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in flow_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # =================================================================
    # State Transitions Table
    # =================================================================
    transitions_section = doc.add_heading('State Transitions', level=2)
    transitions_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    transitions = [
        {
            'from': 'Draft',
            'to': 'Confirmed',
            'method': 'action_confirm',
            'condition': 'ข้อมูลครบถ้วน'
        },
        {
            'from': 'Confirmed',
            'to': 'Paid',
            'method': 'action_pay',
            'condition': 'ชำระเงินสำเร็จ'
        },
        {
            'from': 'Paid',
            'to': 'Assigned',
            'method': 'action_assign_driver',
            'condition': 'มอบหมายคนขับ'
        },
        {
            'from': 'Assigned',
            'to': 'In Progress',
            'method': 'action_start_trip',
            'condition': 'คนขับยืนยัน'
        },
        {
            'from': 'In Progress',
            'to': 'Completed',
            'method': 'action_complete',
            'condition': 'ถึงปลายทาง'
        },
        {
            'from': 'Completed',
            'to': 'Closed',
            'method': 'action_close',
            'condition': 'ยืนยันปิด'
        },
        {
            'from': 'Any',
            'to': 'Cancelled',
            'method': 'action_cancel',
            'condition': 'ยกเลิกโดยผู้ใช้'
        }
    ]
    
    trans_table = doc.add_table(rows=len(transitions) + 1, cols=4)
    trans_table.style = 'Light Grid Accent 1'
    
    # หัวตาราง
    trans_headers = trans_table.rows[0].cells
    header_names = ['From State', 'To State', 'Method', 'Condition']
    
    for i, header_text in enumerate(header_names):
        cell = trans_headers[i]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '366092')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # เพิ่มข้อมูล transitions
    for idx, trans in enumerate(transitions, 1):
        row_cells = trans_table.rows[idx].cells
        row_cells[0].text = trans['from']
        row_cells[1].text = trans['to']
        row_cells[2].text = trans['method']
        row_cells[3].text = trans['condition']
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_page_break()
    
    # =================================================================
    # Implementation Notes
    # =================================================================
    notes_section = doc.add_heading('Implementation Notes', level=2)
    
    notes_items = [
        ('State Definition', 'States ถูกกำหนดใน selection_get() method ของ model'),
        ('Workflow Logic', 'Logic ของการเปลี่ยน state อยู่ใน action methods'),
        ('Validation', 'การ validate ข้อมูลก่อนการเปลี่ยน state'),
        ('Logging', 'บันทึก activity log สำหรับการเปลี่ยน state'),
        ('Notifications', 'ส่งการแจ้งเตือน (Email/Notification) เมื่อเปลี่ยน state')
    ]
    
    for title, description in notes_items:
        p = doc.add_paragraph()
        run = p.add_run(title + ': ')
        run.bold = True
        p.add_run(description)
    
    # บันทึกไฟล์
    output_path = r'C:\Program Files\Odoo 18.0.20251009\server\custom-addons\transport_booking\Workflow_Diagram.docx'
    doc.save(output_path)
    print(f"✓ Workflow Diagram created: {output_path}")
    return output_path

if __name__ == '__main__':
    create_transport_booking_workflow()
