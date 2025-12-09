# -*- coding: utf-8 -*-
"""
Transport Booking System - Workflow Diagram Generator
สร้าง Workflow Diagram เป็น Word Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading_style(doc, text, level=1):
    """เพิ่มหัวข้อ"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading


def add_paragraph_with_format(doc, text, bold=False, italic=False, size=11, color=None):
    """เพิ่มย่อหน้าพร้อมการจัดรูปแบบ"""
    p = doc.add_paragraph(text)
    if bold or italic or size != 11 or color:
        for run in p.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
    return p


def add_box(doc, text, indent_level=0):
    """เพิ่ม Box (กล่อง)"""
    indent = indent_level * 0.5
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # เพิ่มขอบ
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '1')
        border.set(qn('w:color'), '000000')
        pBdr.append(border)
    
    pPr.append(pBdr)
    
    # เพิ่มข้อความ
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_arrow(doc):
    """เพิ่มลูกศร"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("    ↓")
    run.font.size = Pt(14)


def add_table_with_border(doc, rows, cols, data):
    """เพิ่มตาราพร้อมขอบ"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    
    # เพิ่มข้อมูล
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    return table


def create_workflow_document():
    """สร้าง Word document ที่มี workflow diagram"""
    
    doc = Document()
    
    # ================================
    # ส่วนที่ 1: Header
    # ================================
    title = doc.add_heading('Transport Booking System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Workflow Diagram - ไดอะแกรมขั้นตอนการทำงาน', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 2: ภาพรวมระบบ
    # ================================
    doc.add_heading('1. ภาพรวมระบบ (System Overview)', level=1)
    
    overview_text = """
    ระบบ Transport Booking คือระบบจัดการการจองและการขนส่งสินค้า ประกอบด้วย:
    
    • Vehicle Booking: การจองรถขนส่ง
    • Vehicle Tracking: ติดตามตำแหน่งรถ GPS แบบ Real-time
    • Delivery History: บันทึกประวัติการจัดส่ง
    • Delivery Rating: ประเมินความพึงพอใจของลูกค้า
    • Tracking Settings: การตั้งค่าการติดตาม
    """
    
    for line in overview_text.strip().split('\n'):
        if line.strip().startswith('•'):
            doc.add_paragraph(line.strip(), style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line.strip())
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 3: Main Workflow - Vehicle Booking
    # ================================
    doc.add_heading('2. Vehicle Booking - ขั้นตอนหลักของการจอง', level=1)
    
    doc.add_heading('2.1 State Flow (ขั้นตอนสถานะ)', level=2)
    
    # Workflow diagram
    add_box(doc, "[ DRAFT - ร่าง ]")
    add_box(doc, "เอกสารใหม่ที่ยังไม่ยืนยัน")
    add_arrow(doc)
    
    add_box(doc, "[ CONFIRMED - ยืนยันการจอง ]")
    add_box(doc, "เลือกรถและคนขับ → สร้างเลขที่จอง")
    add_box(doc, "เปลี่ยนสถานะรถเป็น 'ติดจอง' (reserved)")
    add_arrow(doc)
    
    add_box(doc, "[ IN_PROGRESS - กำลังขนส่ง ]")
    add_box(doc, "เริ่มงานโดยบันทึกรูปสินค้าก่อนขนส่ง")
    add_box(doc, "เปลี่ยนสถานะรถเป็น 'กำลังจัดส่ง' (in_delivery)")
    add_box(doc, "สร้าง Tracking Record แรก")
    add_arrow(doc)
    
    add_box(doc, "[ DONE - เสร็จสิ้น ]")
    add_box(doc, "บันทึกรูปหลักฐานการส่ง + ลายเซ็นผู้รับ")
    add_box(doc, "คืนสถานะรถเป็น 'พร้อมใช้' (available)")
    add_box(doc, "สร้าง Delivery History Record")
    add_box(doc, "สร้าง Rating Link ให้ลูกค้าประเมิน")
    add_arrow(doc)
    
    add_box(doc, "[ CANCELLED - ยกเลิก ]")
    add_box(doc, "คืนสถานะรถเป็น 'พร้อมใช้' (available)")
    add_box(doc, "ถ้ามีการเริ่มงาน → สร้าง Delivery History")
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 4: Details
    # ================================
    doc.add_heading('2.2 รายละเอียดเพิ่มเติม', level=2)
    
    # Action Details
    doc.add_heading('การกระทำหลัก (Actions)', level=3)
    
    actions = [
        ("action_confirm()", [
            "- ตรวจสอบ: vehicle_id และ driver_id ต้องมีค่า",
            "- สร้างเลขที่จอง: ใช้ ir.sequence 'vehicle.booking'",
            "- อัพเดท Vehicle Status: reserved",
            "- เปลี่ยนสถานะ: draft → confirmed"
        ]),
        ("action_reset_to_draft()", [
            "- คืนสถานะรถเป็น: available",
            "- เปลี่ยนสถานะ: confirmed → draft"
        ]),
        ("action_start()", [
            "- บันทึก actual_pickup_time: fields.Datetime.now()",
            "- อัพเดท Vehicle Status: in_delivery",
            "- เปลี่ยนสถานะ: confirmed → in_progress",
            "- อัพเดท tracking_status: in_transit"
        ]),
        ("start_job_with_photo(photo_base64)", [
            "- บันทึก pickup_photo: Base64 image",
            "- บันทึก actual_pickup_time",
            "- เรียก action_start()",
            "- สร้าง Initial Tracking Record ที่ pickup_location"
        ]),
        ("action_done()", [
            "- บันทึก actual_delivery_time: fields.Datetime.now()",
            "- อัพเดท Vehicle Status: available",
            "- เปลี่ยนสถานะ: in_progress → done",
            "- อัพเดท tracking_status: delivered",
            "- สร้าง Delivery History Record",
            "- สร้าง Rating Link"
        ]),
        ("action_cancel()", [
            "- อัพเดท Vehicle Status: available",
            "- เปลี่ยนสถานะ: confirmed/in_progress → cancelled",
            "- ถ้า actual_pickup_time มีค่า → สร้าง Delivery History"
        ])
    ]
    
    for action_name, details in actions:
        p = doc.add_paragraph(f"{action_name}", style='List Number')
        for detail in details:
            doc.add_paragraph(detail, style='List Bullet 2')
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 5: Tracking Status
    # ================================
    doc.add_heading('2.3 สถานะการติดตาม (Tracking Status)', level=2)
    
    doc.add_paragraph(
        "tracking_status เป็นสถานะเพิ่มเติมสำหรับแสดงความคืบหน้า ของการขนส่ง ไม่เกี่ยวข้องกับเปลี่ยนสถานะหลักของ Booking"
    )
    
    tracking_statuses = [
        ("pending", "รอออกเดินทาง", "ค่าเริ่มต้น เมื่อสร้างการจอง"),
        ("picked_up", "รับสินค้าแล้ว", "เมื่อรับสินค้าที่ต้นทาง"),
        ("in_transit", "อยู่ระหว่างขนส่ง", "ที่อัพเดท auto เมื่อ action_start()"),
        ("near_destination", "ใกล้ถึงปลายทาง", "ใกล้ถึงจุดหมาย"),
        ("delivered", "ส่งถึงแล้ว", "เมื่อ action_done() หรือ action_mark_delivered()")
    ]
    
    for status_code, status_name, description in tracking_statuses:
        p = doc.add_paragraph()
        run = p.add_run(f"{status_code}")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 139)
        p.add_run(f" - {status_name}: {description}")
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 6: Vehicle Tracking
    # ================================
    doc.add_heading('3. Vehicle Tracking - ติดตามตำแหน่ง GPS', level=1)
    
    doc.add_heading('3.1 Flow', level=2)
    
    add_box(doc, "[ Mobile App / GPS Device ]")
    add_box(doc, "บันทึกพิกัด GPS ของรถ")
    add_arrow(doc)
    
    add_box(doc, "[ POST /tracking/create ]")
    add_box(doc, "ส่งข้อมูล GPS ไปยัง Server")
    add_arrow(doc)
    
    add_box(doc, "[ Vehicle Tracking Model ]")
    add_box(doc, "บันทึกข้อมูล Tracking Record")
    add_box(doc, "Fields: booking_id, driver_id, latitude, longitude,")
    add_box(doc, "        timestamp, speed, heading, accuracy, address, notes")
    add_arrow(doc)
    
    add_box(doc, "[ Real-time Map View ]")
    add_box(doc, "แสดงตำแหน่งรถปัจจุบันบนแผนที่")
    add_box(doc, "ดึงข้อมูลจาก get_latest_location()")
    
    doc.add_paragraph()  # ช่องว่าง
    
    doc.add_heading('3.2 Key Functions', level=2)
    
    tracking_functions = [
        ("create_tracking_point(vals)", "สร้าง Tracking Record ใหม่"),
        ("get_latest_location(booking_id)", "ดึงตำแหน่งล่าสุดของ Booking"),
        ("get_tracking_history(booking_id, hours=24)", "ดึงประวัติการเคลื่อนที่ 24 ชั่วโมงย้อนหลัง"),
        ("cleanup_old_tracking(days=30)", "ลบข้อมูล Tracking เก่ากว่า 30 วัน")
    ]
    
    for func_name, description in tracking_functions:
        p = doc.add_paragraph()
        run = p.add_run(f"{func_name}")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 100, 0)
        p.add_run(f"\n{description}")
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 7: Delivery History
    # ================================
    doc.add_heading('4. Delivery History - ประวัติการจัดส่ง', level=1)
    
    doc.add_heading('4.1 Flow', level=2)
    
    add_box(doc, "[ Booking ที่เสร็จสิ้น ]")
    add_box(doc, "state = 'done' หรือ 'cancelled'")
    add_arrow(doc)
    
    add_box(doc, "[ action_done() / action_cancel() ]")
    add_box(doc, "เรียก create_from_booking(booking)")
    add_arrow(doc)
    
    add_box(doc, "[ create_from_booking() ]")
    add_box(doc, "Copy ข้อมูลจาก Booking → Delivery History")
    add_box(doc, "Fields: name, booking_id, partner_id, pickup_location,")
    add_box(doc, "        destination, distance_km, driver_id, vehicle_id,")
    add_box(doc, "        shipping_cost, actual_pickup_time, actual_delivery_time,")
    add_box(doc, "        pickup_photo, delivery_photo, receiver_name, signature")
    add_arrow(doc)
    
    add_box(doc, "[ Delivery History Record ]")
    add_box(doc, "บันทึกถาวร สำหรับ Archive")
    add_box(doc, "Compute: duration_hours = actual_delivery_time - actual_pickup_time")
    
    doc.add_paragraph()  # ช่องว่าง
    
    doc.add_heading('4.2 Key Features', level=2)
    
    history_features = [
        ("Archive", "เก็บข้อมูลการจัดส่งทั้งหมดแบบถาวร"),
        ("Auto Copy", "Copy ข้อมูลจาก Booking โดยอัตโนมัติ"),
        ("Duration", "คำนวณระยะเวลาจัดส่ง (ชั่วโมง) โดยอัตโนมัติ"),
        ("Photo Evidence", "เก็บรูปถ่ายสินค้าและลายเซ็นผู้รับ"),
        ("GPS Data", "เก็บพิกัด GPS ต้นทาง/ปลายทาง"),
    ]
    
    for feature, description in history_features:
        p = doc.add_paragraph(f"{feature}: {description}", style='List Bullet')
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 8: Delivery Rating
    # ================================
    doc.add_heading('5. Delivery Rating - ประเมินความพึงพอใจ', level=1)
    
    doc.add_heading('5.1 Flow', level=2)
    
    add_box(doc, "[ Booking สถานะ 'done' ]")
    add_arrow(doc)
    
    add_box(doc, "[ action_create_rating_link() ]")
    add_box(doc, "สร้าง Delivery Rating Record พร้อม unique token")
    add_box(doc, "Token: UUID - ไม่ซ้ำกันสำหรับแต่ละการจอง")
    add_arrow(doc)
    
    add_box(doc, "[ Generate Public URL ]")
    add_box(doc, "URL: /rating/{rating_token}")
    add_box(doc, "ส่ง Link ให้ลูกค้า (ผ่าน SMS/Email/QR Code)")
    add_arrow(doc)
    
    add_box(doc, "[ Customer Rate ]")
    add_box(doc, "ลูกค้าคลิก Link → เลือกคะแนน (1-5 ดาว)")
    add_box(doc, "เพิ่มความเห็น (optional)")
    add_arrow(doc)
    
    add_box(doc, "[ submit_rating(token, stars, comment) ]")
    add_box(doc, "บันทึกการประเมิน")
    add_box(doc, "state: pending → done")
    add_box(doc, "rating_date: fields.Datetime.now()")
    
    doc.add_paragraph()  # ช่องว่าง
    
    doc.add_heading('5.2 Rating System', level=2)
    
    rating_data = [
        ['คะแนน', 'สัญลักษณ์', 'ความหมาย'],
        ['1', '⭐', 'แย่มาก'],
        ['2', '⭐⭐', 'แย่'],
        ['3', '⭐⭐⭐', 'ปานกลาง'],
        ['4', '⭐⭐⭐⭐', 'ดี'],
        ['5', '⭐⭐⭐⭐⭐', 'ดีมาก'],
    ]
    
    table = add_table_with_border(doc, len(rating_data), 3, rating_data)
    
    # ทำให้ header row เป็น bold
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc.add_paragraph()  # ช่องว่าง
    
    doc.add_heading('5.3 Key Fields', level=2)
    
    rating_fields = [
        ("booking_id", "Many2one", "ตัวเลือก (unique per booking)"),
        ("rating_token", "Char", "UUID - ไม่ซ้ำกัน"),
        ("rating_stars", "Selection", "1-5 ดาว"),
        ("rating_value", "Integer", "แปลง stars เป็นตัวเลข (computed)"),
        ("customer_comment", "Text", "ความเห็นจากลูกค้า"),
        ("rating_date", "Datetime", "วันที่ประเมิน"),
        ("state", "Selection", "pending → done / expired"),
        ("rating_url", "Char", "Public URL (computed)"),
    ]
    
    for field_name, field_type, description in rating_fields:
        p = doc.add_paragraph()
        run = p.add_run(f"{field_name}")
        run.bold = True
        run.font.color.rgb = RGBColor(128, 0, 128)
        p.add_run(f" ({field_type}): {description}")
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 9: Relationships & Data Flow
    # ================================
    doc.add_heading('6. ความสัมพันธ์ระหว่าง Models', level=1)
    
    relationships = [
        ("vehicle.booking", "→", "vehicle.tracking", "One2many", "One Booking มี หลาย Tracking Points"),
        ("vehicle.booking", "→", "delivery.history", "One2one", "One Booking สร้าง One History Record"),
        ("vehicle.booking", "→", "delivery.rating", "One2many", "One Booking มี One Rating Link"),
        ("vehicle.booking", "→", "fleet.vehicle", "Many2one", "Many Bookings สำหรับ One Vehicle"),
        ("vehicle.booking", "→", "vehicle.driver", "Many2one", "Many Bookings สำหรับ One Driver"),
        ("delivery.history", "→", "vehicle.booking", "Many2one", "Reference เพื่อ Archive"),
        ("delivery.rating", "→", "vehicle.booking", "Many2one", "ประเมินสำหรับ Booking"),
    ]
    
    for source, arrow, target, rel_type, description in relationships:
        p = doc.add_paragraph()
        run = p.add_run(source)
        run.bold = True
        p.add_run(f" {arrow} {target} ({rel_type}): {description}")
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 10: Database Relationships Diagram
    # ================================
    doc.add_heading('7. Data Flow Diagram', level=1)
    
    flow_text = """
    Transport Order
         ↓
    Vehicle Booking (Main)
         ├→ fleet.vehicle (Many2one)
         ├→ vehicle.driver (Many2one)
         ├→ res.partner (Many2one) - Customer
         │
         ├→ vehicle.tracking (One2many)
         │  └→ GPS Data + Address
         │
         ├→ delivery.history (created on done)
         │  └→ Archive all booking data
         │
         ├→ delivery.rating (One2many)
         │  └→ Customer review + stars
         │
         └→ vehicle.tracking.history (deprecated, use vehicle.tracking)
    
    Tracking Settings
         └→ res.users (Many2one)
            └→ User GPS tracking preferences
    """
    
    for line in flow_text.strip().split('\n'):
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.2)
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 11: Key Fields ทั้งระบบ
    # ================================
    doc.add_heading('8. Key Fields Reference', level=1)
    
    doc.add_heading('8.1 vehicle.booking - สำคัญ', level=2)
    
    booking_fields = [
        ('name', 'Char', 'Booking Number - สร้างโดย ir.sequence เมื่อ confirm'),
        ('state', 'Selection', 'draft→confirmed→in_progress→done/cancelled'),
        ('tracking_status', 'Selection', 'pending→picked_up→in_transit→near_destination→delivered'),
        ('vehicle_id', 'Many2one', 'รถที่จัดส่ง - ต้องมีและต้อง available'),
        ('driver_id', 'Many2one', 'คนขับ - ต้องมี'),
        ('pickup_photo', 'Binary', 'รูปสินค้าก่อนขนส่ง'),
        ('delivery_photo', 'Binary', 'รูปหลักฐานการส่ง'),
        ('receiver_signature', 'Binary', 'ลายเซ็นผู้รับ'),
        ('actual_pickup_time', 'Datetime', 'เวลารับสินค้าจริง - set เมื่อ action_start()'),
        ('actual_delivery_time', 'Datetime', 'เวลาส่งถึงจริง - set เมื่อ action_done()'),
        ('pickup_latitude/longitude', 'Float', 'GPS พิกัดต้นทาง - auto-geocode'),
        ('destination_latitude/longitude', 'Float', 'GPS พิกัดปลายทาง - auto-geocode'),
        ('tracking_ids', 'One2many', 'GPS Tracking points'),
    ]
    
    for field_name, field_type, description in booking_fields:
        p = doc.add_paragraph()
        run = p.add_run(f"{field_name}")
        run.bold = True
        p.add_run(f" ({field_type}): {description}")
    
    doc.add_paragraph()  # ช่องว่าง
    
    doc.add_heading('8.2 vehicle.tracking - GPS Data', level=2)
    
    tracking_fields = [
        ('booking_id', 'Many2one', 'Reference to Booking - Index'),
        ('driver_id', 'Many2one', 'Driver who submit GPS - Index'),
        ('latitude/longitude', 'Float', 'GPS Coordinates - Required'),
        ('timestamp', 'Datetime', 'When GPS recorded - Index & Default now'),
        ('speed', 'Float', 'Speed in km/h'),
        ('heading', 'Float', 'Direction in degrees'),
        ('accuracy', 'Float', 'GPS accuracy in meters'),
        ('altitude', 'Float', 'Altitude in meters'),
        ('is_moving', 'Boolean', 'Computed: speed > 5 km/h'),
        ('address', 'Char', 'Reverse geocoding address'),
        ('battery_level', 'Float', 'Mobile device battery %'),
    ]
    
    for field_name, field_type, description in tracking_fields:
        p = doc.add_paragraph()
        run = p.add_run(f"{field_name}")
        run.bold = True
        p.add_run(f" ({field_type}): {description}")
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 12: API Endpoints
    # ================================
    doc.add_heading('9. API Endpoints (Controllers)', level=1)
    
    endpoints = [
        ('/tracking/create', 'POST', 'submit GPS tracking data', 'tracking_controller.py'),
        ('/tracking/map/{booking_id}', 'GET', 'show real-time map', 'tracking_controller.py'),
        ('/rating/{token}', 'GET', 'public rating form', 'rating_controller.py'),
        ('/rating/submit', 'POST', 'submit customer rating', 'rating_controller.py'),
        ('/vehicle-map', 'GET', 'show vehicle tracking on map', 'vehicle_tracking_controller.py'),
        ('/map/get-route', 'POST', 'calculate route using Google Maps API', 'map_controller.py'),
    ]
    
    for endpoint, method, description, file in endpoints:
        p = doc.add_paragraph()
        run = p.add_run(f"{method} {endpoint}")
        run.bold = True
        p.add_run(f"\n{description}\nFile: {file}")
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 13: State Transition Logic
    # ================================
    doc.add_heading('10. State Transition Logic', level=1)
    
    doc.add_heading('10.1 Allowed Transitions', level=2)
    
    transitions = [
        ('draft', 'confirmed', 'action_confirm()', ['vehicle_id ต้องมี', 'driver_id ต้องมี', 'สร้างเลขที่จอง']),
        ('confirmed', 'draft', 'action_reset_to_draft()', ['คืนรถเป็น available']),
        ('confirmed', 'in_progress', 'action_start()', ['บันทึกรูปก่อนขนส่ง', 'สร้าง Tracking record']),
        ('confirmed', 'in_progress', 'start_job_with_photo()', ['action_start() + upload photo']),
        ('in_progress', 'done', 'action_done()', ['บันทึกรูปหลักฐาน', 'สร้าง Delivery History', 'สร้าง Rating Link']),
        ('in_progress', 'cancelled', 'action_cancel()', ['สร้าง Delivery History (ถ้า in progress)']),
        ('confirmed', 'cancelled', 'action_cancel()', ['ยกเลิกการจอง']),
    ]
    
    p = doc.add_paragraph('From'.ljust(15) + 'To'.ljust(18) + 'Action'.ljust(25) + 'Requirements/Side Effects')
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.2)
    
    for from_state, to_state, action, effects in transitions:
        p = doc.add_paragraph()
        run = p.add_run(f"{from_state.ljust(15)}{to_state.ljust(18)}{action.ljust(25)}")
        run.font.bold = True
        for effect in effects:
            p.add_run(f"\n  • {effect}")
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ================================
    # ส่วนที่ 14: Summary / Notes
    # ================================
    doc.add_heading('11. สรุปและหมายเหตุ', level=1)
    
    summary_points = [
        ("Main Entry Point", "Transport Order → Vehicle Booking (One2many)"),
        ("Auto Geocoding", "Address → GPS Coordinates (Google Maps API)"),
        ("GPS Tracking", "Real-time via mobile app → vehicle.tracking model"),
        ("Archive", "On completion → create Delivery History record"),
        ("Rating", "After delivery → create Rating Link with unique token"),
        ("Photo Evidence", "pickup_photo + delivery_photo + receiver_signature"),
        ("Duration Calculation", "automatic คำนวณ: actual_delivery_time - actual_pickup_time"),
        ("Vehicle Status", "available → reserved → in_delivery → available (cycle)"),
        ("Public URL", "/rating/{token} - ไม่ต้อง login เพื่อให้คะแนน"),
    ]
    
    for key, value in summary_points:
        p = doc.add_paragraph()
        run = p.add_run(f"{key}")
        run.bold = True
        p.add_run(f": {value}")
    
    doc.add_paragraph()  # ช่องว่าง
    
    doc.add_heading('ไฟล์ที่เกี่ยวข้อง', level=2)
    
    files_list = [
        "models/vehicle_booking.py - Main model และ actions",
        "models/vehicle_tracking.py - GPS tracking data",
        "models/delivery_history.py - Archive completed bookings",
        "models/delivery_rating.py - Customer ratings",
        "models/res_users_settings.py - User tracking preferences",
        "controllers/tracking_controller.py - GPS API endpoints",
        "controllers/rating_controller.py - Rating form endpoints",
        "controllers/vehicle_tracking_controller.py - Map view",
        "controllers/map_controller.py - Google Maps integration",
        "views/vehicle_booking_views.xml - UI forms and lists",
        "data/sequence_data.xml - ir.sequence สำหรับเลขที่จอง",
    ]
    
    for file_path in files_list:
        doc.add_paragraph(file_path, style='List Bullet')
    
    doc.add_paragraph()  # ช่องว่าง
    
    # Save
    output_path = r'C:\Program Files\Odoo 18.0.20251009\server\custom-addons\transport_booking\Transport_Booking_Workflow.docx'
    doc.save(output_path)
    print(f"✅ Document created successfully: {output_path}")


if __name__ == '__main__':
    create_workflow_document()
