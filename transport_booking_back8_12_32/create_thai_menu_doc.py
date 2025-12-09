# -*- coding: utf-8 -*-
"""
Transport Booking System - Thai Menu Structure Documentation
เอกสารรวมเมนูภาษาไทยทั้งหมด
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_thai_menu_document():
    """สร้าง Document ที่รวมเมนูภาษาไทยทั้งหมด"""
    
    doc = Document()
    
    # Header
    title = doc.add_heading('Transport Booking System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('โครงสร้างเมนูภาษาไทย (Thai Menu Structure)', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ===================================
    # Main Menu
    # ===================================
    doc.add_heading('1. เมนูหลัก (Main Menu)', level=1)
    
    main_menu_table = doc.add_table(rows=1, cols=3)
    main_menu_table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = main_menu_table.rows[0].cells
    header_cells[0].text = 'Menu ID'
    header_cells[1].text = 'ชื่อเมนู (Thai)'
    header_cells[2].text = 'ไอคอน'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Main menu item
    row = main_menu_table.add_row().cells
    row[0].text = 'menu_vehicle_booking_root'
    row[1].text = '🚚 จองคิวรถขนส่ง'
    row[2].text = 'transport_booking/static/description/icon.png'
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ===================================
    # Sub Menus
    # ===================================
    doc.add_heading('2. เมนูย่อย (Sub Menus)', level=1)
    
    submenu_data = [
        ['Menu ID', 'ชื่อเมนู (Thai)', 'Sequence', 'Description'],
        ['menu_vehicle_booking', '📋 จองคิวรถขนส่ง', '10', 'สร้างและจัดการการจอง'],
        ['menu_vehicle_tracking', '📍 ติดตามการขนส่ง', 'ภายในการจอง', 'ดูตำแหน่ง GPS (Tab)'],
        ['menu_delivery_history', '📜 ประวัติการจัดส่ง', 'ภายในการจอง', 'ข้อมูลการจัดส่งเสร็จ (Tab)'],
        ['menu_delivery_rating', '⭐ ประเมินความพึงพอใจ', 'ภายในการจอง', 'Link ประเมินลูกค้า (Tab)'],
        ['menu_tracking_settings', '⚙️ ตั้งค่าการติดตาม', '50', 'ตั้งค่าการติดตาม (ระบบ)'],
    ]
    
    submenu_table = doc.add_table(rows=len(submenu_data), cols=4)
    submenu_table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(submenu_data):
        for j, cell_data in enumerate(row_data):
            submenu_table.rows[i].cells[j].text = cell_data
            if i == 0:  # Header row
                for paragraph in submenu_table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ===================================
    # Tab-based Menu (within Vehicle Booking)
    # ===================================
    doc.add_heading('3. แท็บในฟอร์มการจอง (Tabs in Booking Form)', level=1)
    
    doc.add_paragraph(
        'ระบบใช้งานแท็บแทนเมนูย่อยเพื่อให้ผู้ใช้เข้าถึงข้อมูลได้ง่ายขึ้น:'
    )
    
    tabs_data = [
        ['ลำดับ', 'ชื่อแท็บ (Thai)', 'ไฟล์'],
        ['1', '📦 ข้อมูลคำสั่งขนส่ง', 'vehicle_booking_views.xml'],
        ['2', '🚚 จัดสรรรถและคนขับ', 'vehicle_booking_views.xml'],
        ['3', '📍 สถานที่และเส้นทาง', 'vehicle_booking_views.xml'],
        ['4', '📸 ข้อมูลการรับและส่ง', 'vehicle_booking_views.xml'],
        ['5', '📍 GPS Tracking', 'vehicle_booking_views.xml'],
        ['6', '⭐ ประเมินความพึงพอใจ', 'vehicle_booking_views.xml'],
        ['7', '📝 หมายเหตุ', 'vehicle_booking_views.xml'],
    ]
    
    tabs_table = doc.add_table(rows=len(tabs_data), cols=3)
    tabs_table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(tabs_data):
        for j, cell_data in enumerate(row_data):
            tabs_table.rows[i].cells[j].text = cell_data
            if i == 0:
                for paragraph in tabs_table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ===================================
    # Views List
    # ===================================
    doc.add_heading('4. Views ทั้งหมด (ภาษาไทย)', level=1)
    
    views_list = [
        ('vehicle_booking_views.xml', [
            '📋 List View: "จองคิวรถขนส่ง"',
            '📝 Form View: "จองคิวรถขนส่ง"',
            '🔍 Search View: "ค้นหาการจอง"',
            '⚡ Kanban View: "การจองและสถานะ"'
        ]),
        ('vehicle_tracking_views.xml', [
            '📍 List View: "ติดตามตำแหน่งรถ"',
            '📋 Form View: "รายละเอียดการติดตาม"',
            '📊 Kanban View: "การติดตาม (Active)"',
            '🔍 Search View: "ค้นหาการติดตาม"'
        ]),
        ('delivery_history_views.xml', [
            '📜 List View: "ประวัติการจัดส่ง"',
            '📝 Form View: "ประวัติการจัดส่ง"',
            '📊 Kanban View: "ประวัติ (Card)"',
            '🔍 Search View: "ค้นหาประวัติการจัดส่ง"'
        ]),
        ('delivery_rating_views.xml', [
            '⭐ List View: "การประเมินความพึงพอใจ"',
            '📝 Form View: "การประเมินความพึงพอใจ"'
        ]),
        ('res_users_settings_views.xml', [
            '⚙️ Form View: "⚙️ ตั้งค่าการติดตามรถ"',
            '📋 List View: "การตั้งค่าผู้ใช้"'
        ]),
        ('tracking_settings_views.xml', [
            '⚙️ Form View: "ตั้งค่าการติดตาม"'
        ]),
    ]
    
    for file_name, views in views_list:
        p = doc.add_heading(f'📄 {file_name}', level=2)
        for view in views:
            doc.add_paragraph(view, style='List Bullet')
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ===================================
    # Button & Action Labels
    # ===================================
    doc.add_heading('5. ป้ายกำกับปุ่มและการกระทำ (Buttons & Actions)', level=1)
    
    buttons_data = [
        ['Type', 'Label (Thai)', 'Description'],
        ['Button', '✅ ยืนยันการจอง', 'action_confirm()'],
        ['Button', '🔄 รีเซ็ตเป็นร่าง', 'action_reset_to_draft()'],
        ['Button', '🚚 เริ่มขนส่ง', 'action_start()'],
        ['Button', '✔️ เสร็จสิ้น', 'action_done()'],
        ['Button', '❌ ยกเลิก', 'action_cancel()'],
        ['Button', '🔗 ส่ง Link ประเมิน', 'action_send_rating_link()'],
        ['Button', '📝 สร้าง Link ประเมินใหม่', 'action_create_rating_link()'],
        ['Button', '🗺️ ดูแผนที่', 'action_view_map()'],
        ['Button', '📍 ดูการติดตาม', 'action_view_tracking()'],
        ['Smart Button', '📍 ตำแหน่ง GPS', 'tracking_count'],
    ]
    
    buttons_table = doc.add_table(rows=len(buttons_data), cols=3)
    buttons_table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(buttons_data):
        for j, cell_data in enumerate(row_data):
            buttons_table.rows[i].cells[j].text = cell_data
            if i == 0:
                for paragraph in buttons_table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ===================================
    # Status & Badges
    # ===================================
    doc.add_heading('6. สถานะและป้ายกำกับ (Status & Badges)', level=1)
    
    status_data = [
        ['State Code', 'Label (Thai)', 'Color'],
        ['draft', '📝 ร่าง', 'info'],
        ['confirmed', '✅ ยืนยันการจอง', 'primary'],
        ['in_progress', '🚚 กำลังขนส่ง', 'warning'],
        ['done', '✔️ เสร็จสิ้น', 'success'],
        ['cancelled', '❌ ยกเลิก', 'danger'],
        ['', '', ''],
        ['Tracking Status', 'Label (Thai)', ''],
        ['pending', 'รอออกเดินทาง', ''],
        ['picked_up', 'รับสินค้าแล้ว', ''],
        ['in_transit', 'อยู่ระหว่างขนส่ง', ''],
        ['near_destination', 'ใกล้ถึงปลายทาง', ''],
        ['delivered', 'ส่งถึงแล้ว', ''],
    ]
    
    status_table = doc.add_table(rows=len(status_data), cols=3)
    status_table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(status_data):
        for j, cell_data in enumerate(row_data):
            status_table.rows[i].cells[j].text = cell_data
            if i in [0, 7]:  # Header rows
                for paragraph in status_table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ===================================
    # Form Groups & Sections
    # ===================================
    doc.add_heading('7. หัวข้อฟอร์มหลัก (Form Groups)', level=1)
    
    groups_list = [
        ('Vehicle Booking Form', [
            '📦 ข้อมูลคำสั่งขนส่ง',
            '🚚 จัดสรรรถและคนขับ',
            '📍 สถานที่',
            '💰 ค่าใช้จ่าย',
            '📸 หลักฐาน',
            '👤 ข้อมูลผู้รับ',
            '📍 GPS Tracking',
            '📝 บันทึกการติดตาม',
            '📊 สถิติการประเมิน',
            '🔗 สร้าง Link ประเมิน',
            '⭐ ประวัติการประเมินทั้งหมด',
            '📝 หมายเหตุ'
        ]),
        ('Delivery History Form', [
            'ข้อมูลทั่วไป',
            'พนักงานและรถ',
            'เส้นทาง',
            'เวลา',
            '📸 หลักฐาน',
            '📍 พิกัด GPS',
            '📝 หมายเหตุ'
        ]),
        ('Tracking Settings Form', [
            'การติดตามตำแหน่ง',
            'การแจ้งเตือน',
            'การแสดงผลแผนที่',
            'การบันทึกข้อมูล'
        ]),
    ]
    
    for form_name, groups in groups_list:
        p = doc.add_heading(form_name, level=2)
        for group in groups:
            doc.add_paragraph(group, style='List Bullet')
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ===================================
    # Search Filters
    # ===================================
    doc.add_heading('8. ตัวกรองค้นหา (Search Filters)', level=1)
    
    search_list = [
        ('Vehicle Booking Search', [
            'ค้นหาตามเลขที่จอง',
            'ค้นหาตามลูกค้า',
            'ค้นหาตามรถ',
            'ค้นหาตามคนขับ',
            'ตัวกรอง: งานร่าง',
            'ตัวกรอง: งานที่ยืนยัน',
            'ตัวกรอง: งานที่กำลังทำ',
            'ตัวกรอง: งานเสร็จสิ้น',
            'จัดกลุ่มตามรถ',
            'จัดกลุ่มตามคนขับ',
            'จัดกลุ่มตามลูกค้า',
            'จัดกลุ่มตามสถานะ'
        ]),
        ('Tracking Search', [
            'วันนี้',
            'สัปดาห์นี้',
            'กำลังเคลื่อนที่',
            'งานที่กำลังทำ',
            'จัดกลุ่มตามการจอง',
            'จัดกลุ่มตามคนขับ',
            'จัดกลุ่มตามรถ',
            'จัดกลุ่มตามวันที่'
        ]),
        ('Delivery History Search', [
            '✅ เสร็จสิ้น',
            '❌ ยกเลิก',
            'วันนี้',
            'สัปดาห์นี้',
            'เดือนนี้',
            'จัดกลุ่มตามคนขับ',
            'จัดกลุ่มตามรถ',
            'จัดกลุ่มตามลูกค้า',
            'จัดกลุ่มตามวันที่',
            'จัดกลุ่มตามสถานะ'
        ]),
    ]
    
    for search_name, filters in search_list:
        p = doc.add_heading(search_name, level=2)
        for filter_item in filters:
            doc.add_paragraph(filter_item, style='List Bullet')
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ===================================
    # Summary
    # ===================================
    doc.add_heading('9. สรุป', level=1)
    
    summary = [
        ('✅ เมนูภาษาไทยทั้งหมด', 'ระบบใช้ภาษาไทยเต็มที่ในเมนูและ UI'),
        ('📍 ตำแหน่งเมนู', 'เมนูหลัก: "จองคิวรถขนส่ง" + เมนูย่อยแบบแท็บ'),
        ('🎨 ไอคอนและ Emoji', 'ใช้ Emoji เพื่อให้ UI มีความน่าสนใจและง่ายในการบ่งชี้'),
        ('📱 Mobile Friendly', 'เมนูและ UI สามารถใช้งานบนมือถือได้ดี'),
        ('🌐 Multi-language Ready', 'สามารถเพิ่มภาษาอื่นได้ง่ายในอนาคต'),
        ('📊 Data Entry', 'ทุก Field มีป้ายกำกับภาษาไทย'),
        ('🔍 Search & Filter', 'ตัวกรองและการค้นหาทั้งหมดใช้ภาษาไทย'),
    ]
    
    for point, description in summary:
        p = doc.add_paragraph()
        run = p.add_run(f"{point}")
        run.bold = True
        p.add_run(f": {description}")
    
    doc.add_paragraph()  # ช่องว่าง
    
    # ===================================
    # Files Modified
    # ===================================
    doc.add_heading('10. ไฟล์ที่แก้ไข', level=1)
    
    doc.add_paragraph('ไฟล์ View ที่ตรวจสอบและแก้ไขเป็นภาษาไทยเต็มที่:', style='List Bullet')
    
    files_modified = [
        'views/vehicle_booking_views.xml - ✅ ภาษาไทยทั้งหมด',
        'views/vehicle_tracking_views.xml - ✅ แก้ไขให้เป็นภาษาไทย',
        'views/delivery_history_views.xml - ✅ ภาษาไทยทั้งหมด',
        'views/delivery_rating_views.xml - ✅ ภาษาไทยทั้งหมด',
        'views/res_users_settings_views.xml - ✅ ภาษาไทยทั้งหมด',
        'views/tracking_settings_views.xml - ✅ แก้ไขให้เป็นภาษาไทย',
        'views/rating_templates.xml - ✅ ภาษาไทยทั้งหมด',
        'views/tracking_map_template.xml - ✅ ภาษาไทยทั้งหมด',
        '__manifest__.py - ✅ ภาษาไทยทั้งหมด',
    ]
    
    for file_path in files_modified:
        doc.add_paragraph(file_path, style='List Bullet 2')
    
    doc.add_paragraph()  # ช่องว่าง
    
    # Save
    output_path = r'C:\Program Files\Odoo 18.0.20251009\server\custom-addons\transport_booking\Thai_Menu_Documentation.docx'
    doc.save(output_path)
    print(f"✅ Thai Menu Documentation created: {output_path}")


if __name__ == '__main__':
    create_thai_menu_document()
